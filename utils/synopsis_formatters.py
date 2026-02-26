"""
Форматтеры для генерации синопсиса в различных форматах
"""
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)

# Условный импорт docx (может быть не установлен)
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_docx_synopsis(synopsis_data: Dict[str, Any], output_path: str) -> str:
    """
    Генерирует полный DOCX синопсис со всеми секциями протокола
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx не установлен. Установите: py -m pip install python-docx")
    
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('СИНОПСИС ПРОТОКОЛА', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Исследование биоэквивалентности', 1)
    doc.add_paragraph(f"Препарат: {synopsis_data['inn']}")
    doc.add_paragraph(f"Форма выпуска: {synopsis_data['dosage_form']}")
    doc.add_paragraph(f"Дозировка: {synopsis_data['dosage']}")
    doc.add_paragraph(f"Способ введения: {synopsis_data['administration_mode']}")
    doc.add_paragraph(f"Дата генерации: {synopsis_data['generated_date']}")
    doc.add_paragraph('')
    
    # 1. OBJECTIVE
    doc.add_heading('1. ЦЕЛЬ ИССЛЕДОВАНИЯ', 1)
    objective = synopsis_data.get('objective', {})
    doc.add_paragraph(f"Основная цель: {objective.get('primary', 'N/A')}")
    doc.add_paragraph('Вторичные цели:')
    for sec_obj in objective.get('secondary', []):
        doc.add_paragraph(sec_obj, style='List Bullet')
    doc.add_paragraph('')
    
    # 2. STUDY DESIGN
    doc.add_heading('2. ДИЗАЙН ИССЛЕДОВАНИЯ', 1)
    design = synopsis_data.get('study_design', {})
    doc.add_paragraph(f"Тип дизайна: {design.get('design', 'N/A')}")
    doc.add_paragraph(f"Тип исследования: {design.get('type', 'N/A')}")
    doc.add_paragraph(f"Количество периодов: {design.get('periods', 'N/A')}")
    doc.add_paragraph(f"Washout период: {design.get('washout_duration', 'N/A')}")
    doc.add_paragraph(f"Обоснование: {design.get('rationale', 'N/A')}")
    doc.add_paragraph('')
    
    # 3. POPULATION
    doc.add_heading('3. ИССЛЕДУЕМАЯ ПОПУЛЯЦИЯ', 1)
    population = synopsis_data.get('population', {})
    doc.add_paragraph(f"Тип: {population.get('type', 'N/A')}")
    doc.add_paragraph(f"Возраст: {population.get('age_range', 'N/A')}")
    doc.add_paragraph(f"Пол: {population.get('gender', 'N/A')}")
    doc.add_paragraph(f"Общее количество субъектов: {population.get('total_subjects', 'N/A')}")
    doc.add_paragraph('')
    
    # 4. INCLUSION CRITERIA
    doc.add_heading('4. КРИТЕРИИ ВКЛЮЧЕНИЯ', 1)
    for criterion in synopsis_data.get('inclusion_criteria', []):
        doc.add_paragraph(criterion, style='List Bullet')
    doc.add_paragraph('')
    
    # 5. EXCLUSION CRITERIA
    doc.add_heading('5. КРИТЕРИИ ИСКЛЮЧЕНИЯ', 1)
    for criterion in synopsis_data.get('exclusion_criteria', []):
        doc.add_paragraph(criterion, style='List Bullet')
    doc.add_paragraph('')
    
    # 6. PK SAMPLING
    doc.add_heading('6. СХЕМА ЗАБОРА ПРОБ ДЛЯ ФАРМАКОКИНЕТИКИ', 1)
    pk_sampling = synopsis_data.get('pk_sampling', {})
    doc.add_paragraph('Точки забора проб:')
    for point in pk_sampling.get('scheme', []):
        doc.add_paragraph(point, style='List Bullet')
    doc.add_paragraph(f"Общее количество образцов: {pk_sampling.get('total_samples', 'N/A')}")
    doc.add_paragraph(f"Объем пробы: {pk_sampling.get('sample_volume', 'N/A')}")
    doc.add_paragraph('')
    
    # 7. ENDPOINTS
    doc.add_heading('7. КРИТЕРИИ ОЦЕНКИ (ENDPOINTS)', 1)
    endpoints = synopsis_data.get('endpoints', {})
    doc.add_paragraph('Основные критерии:')
    for endpoint in endpoints.get('primary', []):
        if isinstance(endpoint, dict):
            doc.add_paragraph(f"{endpoint.get('parameter', 'N/A')}: {endpoint.get('description', 'N/A')}")
            doc.add_paragraph(f"Критерий: {endpoint.get('criterion', 'N/A')}", style='List Bullet 2')
        else:
            doc.add_paragraph(str(endpoint), style='List Bullet')
    doc.add_paragraph('Вторичные критерии:')
    for endpoint in endpoints.get('secondary', []):
        doc.add_paragraph(endpoint, style='List Bullet')
    doc.add_paragraph('')
    
    # 8. BIOANALYSIS
    doc.add_heading('8. БИОАНАЛИТИЧЕСКИЙ МЕТОД', 1)
    bioanalysis = synopsis_data.get('bioanalysis', {})
    doc.add_paragraph(f"Метод: {bioanalysis.get('method', 'N/A')}")
    doc.add_paragraph(f"Валидация: {bioanalysis.get('validation', 'N/A')}")
    doc.add_paragraph('Параметры валидации:')
    params = bioanalysis.get('parameters', {})
    for param, value in params.items():
        doc.add_paragraph(f"{param}: {value}", style='List Bullet')
    doc.add_paragraph('')
    
    # 9. SAFETY
    doc.add_heading('9. БЕЗОПАСНОСТЬ', 1)
    safety = synopsis_data.get('safety', {})
    doc.add_paragraph('Мониторинг безопасности:')
    for item in safety.get('monitoring', []):
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('')
    
    # 10. STATISTICAL METHODS
    doc.add_heading('10. СТАТИСТИЧЕСКАЯ МЕТОДОЛОГИЯ', 1)
    stats = synopsis_data.get('statistical_methods', {})
    doc.add_paragraph(f"Популяция анализа: {stats.get('analysis_population', 'N/A')}")
    doc.add_paragraph(f"Метод: {stats.get('method', 'N/A')}")
    doc.add_paragraph(f"Преобразование: {stats.get('transformation', 'N/A')}")
    doc.add_paragraph(f"Модель: {stats.get('model', 'N/A')}")
    doc.add_paragraph(f"Мощность: {stats.get('power', 'N/A')}")
    doc.add_paragraph(f"Критерий биоэквивалентности: {stats.get('significance', 'N/A')}")
    doc.add_paragraph('')
    
    # 11. PK PARAMETERS
    doc.add_heading('11. ФАРМАКОКИНЕТИЧЕСКИЕ ПАРАМЕТРЫ', 1)
    pk_params = synopsis_data.get('pk_parameters', {})
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Параметр'
    headers[1].text = 'Значение'
    headers[2].text = 'Единица'
    
    params_list = [
        ('Cmax', pk_params.get('cmax', {})),
        ('AUC', pk_params.get('auc', {})),
        ('Tmax', pk_params.get('tmax', {})),
        ('T½', pk_params.get('t_half', {}))
    ]
    
    for i, (param_name, param_data) in enumerate(params_list, start=1):
        row = table.rows[i].cells
        row[0].text = param_name
        if isinstance(param_data, dict):
            row[1].text = str(param_data.get('value', 'N/A'))
            row[2].text = param_data.get('unit', 'N/A')
        else:
            row[1].text = 'N/A'
            row[2].text = 'N/A'
    
    doc.add_paragraph('')
    
    # 12. SAMPLE SIZE
    doc.add_heading('12. РАЗМЕР ВЫБОРКИ', 1)
    sample = synopsis_data.get('sample_size', {})
    doc.add_paragraph(f"CVintra: {sample.get('cvintra', 'N/A')}%")
    doc.add_paragraph(f"Базовый размер: {sample.get('base_sample_size', 'N/A')}")
    doc.add_paragraph(f"Итоговый размер: {sample.get('final_sample_size', 'N/A')}")
    if sample.get('calculation_steps'):
        doc.add_paragraph('Расчет:')
        for step in sample['calculation_steps']:
            doc.add_paragraph(step, style='List Number')
    doc.add_paragraph('')
    
    # 13. REGULATORY COMPLIANCE
    doc.add_heading('13. РЕГУЛЯТОРНОЕ СООТВЕТСТВИЕ', 1)
    regulatory = synopsis_data.get('regulatory_compliance', {})
    if not regulatory:
        # Fallback на старый формат
        regulatory = synopsis_data.get('regulatory_check', {})
    
    for reg_name, reg_data in regulatory.items():
        if isinstance(reg_data, dict):
            status = '✓ Соответствует' if reg_data.get('compliant') else '✗ Не проверено'
            doc.add_paragraph(f"{reg_name.upper()}: {status}")
            req = reg_data.get('requirements', '')
            if req:
                doc.add_paragraph(req, style='List Bullet 2')
        else:
            doc.add_paragraph(f"{reg_name.upper()}: {reg_data}")
    doc.add_paragraph('')
    
    # Сохранение
    try:
        doc.save(output_path)
        logger.info(f"DOCX файл сохранен: {output_path}")
    except Exception as e:
        logger.error(f"Ошибка сохранения DOCX: {e}", exc_info=True)
        raise
    
    return output_path


def generate_markdown_synopsis(synopsis_data: Dict[str, Any]) -> str:
    """
    Генерирует полный Markdown синопсис со всеми секциями протокола
    """
    md = f"""# {synopsis_data['title']}

## Исследуемый препарат

- **МНН:** {synopsis_data['inn']}
- **Форма выпуска:** {synopsis_data['dosage_form']}
- **Дозировка:** {synopsis_data['dosage']}
- **Способ введения:** {synopsis_data['administration_mode']}
- **Дата генерации:** {synopsis_data['generated_date']}

---

## 1. ЦЕЛЬ ИССЛЕДОВАНИЯ

### Основная цель
{synopsis_data.get('objective', {}).get('primary', 'N/A')}

### Вторичные цели
"""
    
    for sec_obj in synopsis_data.get('objective', {}).get('secondary', []):
        md += f"\n- {sec_obj}"
    
    design = synopsis_data.get('study_design', {})
    md += f"""

---

## 2. ДИЗАЙН ИССЛЕДОВАНИЯ

| Параметр | Значение |
|----------|----------|
| **Дизайн** | {design.get('design', 'N/A')} |
| **Тип исследования** | {design.get('type', 'N/A')} |
| **Количество периодов** | {design.get('periods', 'N/A')} |
| **Washout период** | {design.get('washout_duration', 'N/A')} |
| **Обоснование** | {design.get('rationale', 'N/A')} |

---

## 3. ИССЛЕДУЕМАЯ ПОПУЛЯЦИЯ

- **Тип:** {synopsis_data.get('population', {}).get('type', 'N/A')}
- **Возраст:** {synopsis_data.get('population', {}).get('age_range', 'N/A')}
- **Пол:** {synopsis_data.get('population', {}).get('gender', 'N/A')}
- **Общее количество субъектов:** {synopsis_data.get('population', {}).get('total_subjects', 'N/A')}

---

## 4. КРИТЕРИИ ВКЛЮЧЕНИЯ

"""
    
    for criterion in synopsis_data.get('inclusion_criteria', []):
        md += f"\n- {criterion}"
    
    md += "\n\n---\n\n## 5. КРИТЕРИИ ИСКЛЮЧЕНИЯ\n\n"
    
    for criterion in synopsis_data.get('exclusion_criteria', []):
        md += f"\n- {criterion}"
    
    pk_sampling = synopsis_data.get('pk_sampling', {})
    md += f"""

---

## 6. СХЕМА ЗАБОРА ПРОБ ДЛЯ ФАРМАКОКИНЕТИКИ

"""
    
    for point in pk_sampling.get('scheme', []):
        md += f"\n- {point}"
    
    md += f"\n\n- **Общее количество образцов:** {pk_sampling.get('total_samples', 'N/A')}"
    md += f"\n- **Объем пробы:** {pk_sampling.get('sample_volume', 'N/A')}"
    
    endpoints = synopsis_data.get('endpoints', {})
    md += "\n\n---\n\n## 7. КРИТЕРИИ ОЦЕНКИ (ENDPOINTS)\n\n### Основные критерии\n\n"
    
    for endpoint in endpoints.get('primary', []):
        if isinstance(endpoint, dict):
            md += f"- **{endpoint.get('parameter', 'N/A')}:** {endpoint.get('description', 'N/A')}\n"
            md += f"  - Критерий: {endpoint.get('criterion', 'N/A')}\n"
        else:
            md += f"- {endpoint}\n"
    
    md += "\n### Вторичные критерии\n\n"
    for endpoint in endpoints.get('secondary', []):
        md += f"- {endpoint}\n"
    
    bioanalysis = synopsis_data.get('bioanalysis', {})
    md += f"""

---

## 8. БИОАНАЛИТИЧЕСКИЙ МЕТОД

- **Метод:** {bioanalysis.get('method', 'N/A')}
- **Валидация:** {bioanalysis.get('validation', 'N/A')}

### Параметры валидации

"""
    
    for param, value in bioanalysis.get('parameters', {}).items():
        md += f"- **{param}:** {value}\n"
    
    safety = synopsis_data.get('safety', {})
    md += "\n---\n\n## 9. БЕЗОПАСНОСТЬ\n\n### Мониторинг безопасности\n\n"
    
    for item in safety.get('monitoring', []):
        md += f"- {item}\n"
    
    stats = synopsis_data.get('statistical_methods', {})
    md += f"""

---

## 10. СТАТИСТИЧЕСКАЯ МЕТОДОЛОГИЯ

- **Популяция анализа:** {stats.get('analysis_population', 'N/A')}
- **Метод:** {stats.get('method', 'N/A')}
- **Преобразование:** {stats.get('transformation', 'N/A')}
- **Модель:** {stats.get('model', 'N/A')}
- **Мощность:** {stats.get('power', 'N/A')}
- **Критерий биоэквивалентности:** {stats.get('significance', 'N/A')}

---

## 11. ФАРМАКОКИНЕТИЧЕСКИЕ ПАРАМЕТРЫ

| Параметр | Значение | Единица |
|----------|----------|---------|
"""
    
    pk_params = synopsis_data.get('pk_parameters', {})
    params_list = [
        ('Cmax', pk_params.get('cmax', {})),
        ('AUC', pk_params.get('auc', {})),
        ('Tmax', pk_params.get('tmax', {})),
        ('T½', pk_params.get('t_half', {}))
    ]
    
    for param_name, param_data in params_list:
        if isinstance(param_data, dict):
            value = str(param_data.get('value', 'N/A'))
            unit = param_data.get('unit', 'N/A')
        else:
            value = 'N/A'
            unit = 'N/A'
        md += f"| {param_name} | {value} | {unit} |\n"
    
    sample = synopsis_data.get('sample_size', {})
    md += f"""

---

## 12. РАЗМЕР ВЫБОРКИ

| Параметр | Значение |
|----------|----------|
| CVintra | {sample.get('cvintra', 'N/A')}% |
| Базовый размер выборки | {sample.get('base_sample_size', 'N/A')} участников |
| Процент выбытия | {sample.get('dropout_rate', 'N/A')}% |
| **Итоговый размер выборки** | **{sample.get('final_sample_size', 'N/A')} участников** |

### Этапы расчета:

"""
    
    for step in sample.get('calculation_steps', []):
        md += f"\n{step}"
    
    regulatory = synopsis_data.get('regulatory_compliance', {})
    md += "\n\n---\n\n## 13. РЕГУЛЯТОРНОЕ СООТВЕТСТВИЕ\n\n"
    
    for reg_name, reg_data in regulatory.items():
        if isinstance(reg_data, dict):
            status = "✓ **Соответствует**" if reg_data.get('compliant') else "✗ **Не соответствует**"
            requirements = reg_data.get('requirements', 'Информация недоступна')
            md += f"\n### {reg_name.upper()}: {status}\n{requirements}\n"
    
    md += "\n---\n\n*Документ автоматически сгенерирован системой BE Study Design AI Assistant*\n"
    
    return md