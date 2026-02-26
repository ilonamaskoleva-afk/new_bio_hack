from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class SynopsisGenerator:
    """
    Генерация синопсиса протокола BE исследования
    """
    
    @staticmethod
    def generate_docx(data: dict, output_path: str = "synopsis.docx") -> str:
        """
        Генерация Word документа
        """
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('СИНОПСИС ПРОТОКОЛА', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('Исследование биоэквивалентности', 1)
        doc.add_paragraph(f"Препарат: {data.get('inn', 'N/A')}")
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_paragraph('')
        
        # 1. Обоснование дизайна
        doc.add_heading('1. ОБОСНОВАНИЕ ДИЗАЙНА ИССЛЕДОВАНИЯ', 1)
        doc.add_paragraph(data.get('design_rationale', 'N/A'))
        doc.add_paragraph('')
        
        # 2. Цели и задачи
        doc.add_heading('2. ЦЕЛИ И ЗАДАЧИ', 1)
        doc.add_paragraph('Цель: Оценка биоэквивалентности исследуемого препарата и референтного препарата.')
        doc.add_paragraph('')
        
        # 3. Дизайн исследования
        doc.add_heading('3. ДИЗАЙН ИССЛЕДОВАНИЯ', 1)
        design = data.get('study_design', {})
        doc.add_paragraph(f"Тип дизайна: {design.get('design', 'N/A')}")
        doc.add_paragraph(f"Количество периодов: {design.get('periods', 'N/A')}")
        doc.add_paragraph(f"Washout период: {design.get('washout_duration', 'N/A')}")
        doc.add_paragraph('')
        
        # 4. Популяция
        doc.add_heading('4. ИССЛЕДУЕМАЯ ПОПУЛЯЦИЯ', 1)
        doc.add_paragraph('Здоровые добровольцы мужского и женского пола в возрасте 18-45 лет.')
        doc.add_paragraph('')
        
        # 5. Размер выборки
        doc.add_heading('5. РАЗМЕР ВЫБОРКИ', 1)
        sample_size = data.get('sample_size', {})
        doc.add_paragraph(f"CVintra: {sample_size.get('cvintra', 'N/A')}%")
        doc.add_paragraph(f"Базовый размер: {sample_size.get('base_sample_size', 'N/A')}")
        doc.add_paragraph(f"С учетом drop-out ({sample_size.get('dropout_rate', 'N/A')}%): {sample_size.get('final_sample_size', 'N/A')}")
        
        # Формула
        doc.add_paragraph('')
        doc.add_paragraph('Расчет:')
        for step in sample_size.get('steps', []):
            doc.add_paragraph(step, style='List Number')
        doc.add_paragraph('')
        
        # 6. PK параметры
        doc.add_heading('6. ФАРМАКОКИНЕТИЧЕСКИЕ ПАРАМЕТРЫ', 1)
        pk_params = data.get('pk_parameters', {})
        
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'
        
        headers = table.rows[0].cells
        headers[0].text = 'Параметр'
        headers[1].text = 'Описание'
        headers[2].text = 'Единица'
        
        params_list = [
            ('Cmax', 'Максимальная концентрация в плазме', 'нг/мл'),
            ('AUC', 'Площадь под кривой концентрация-время', 'нг·ч/мл'),
            ('Tmax', 'Время достижения Cmax', 'ч'),
            ('T½', 'Период полувыведения', 'ч'),
            ('CVintra', 'Внутрисубъектная вариабельность', '%')
        ]
        
        for i, (param, desc, unit) in enumerate(params_list, start=1):
            row = table.rows[i].cells
            row[0].text = param
            row[1].text = desc
            row[2].text = unit
        
        doc.add_paragraph('')
        
        # 7. Статистическая методология
        doc.add_heading('7. СТАТИСТИЧЕСКАЯ МЕТОДОЛОГИЯ', 1)
        doc.add_paragraph('Критерий биоэквивалентности: 90% доверительный интервал для отношения геометрических средних Cmax и AUC должен находиться в пределах 80.00% - 125.00%.')
        doc.add_paragraph('')
        
        # 8. Регуляторное соответствие
        doc.add_heading('8. СООТВЕТСТВИЕ РЕГУЛЯТОРНЫМ ТРЕБОВАНИЯМ', 1)
        regulatory = data.get('regulatory_check', {})
        doc.add_paragraph(f"Решение № 85 (РФ): {'✓ Соответствует' if regulatory.get('russia_decision_85', {}).get('compliant', False) else '✗ Не соответствует'}")
        doc.add_paragraph(f"EMA Guidelines: {'✓ Соответствует' if regulatory.get('ema', {}).get('compliant', False) else '✗ Не соответствует'}")
        doc.add_paragraph(f"FDA Guidance: {'✓ Соответствует' if regulatory.get('fda', {}).get('compliant', False) else '✗ Не соответствует'}")
        doc.add_paragraph('')
        
        # 9. Библиография
        doc.add_heading('9. БИБЛИОГРАФИЧЕСКИЙ СПИСОК', 1)
        sources = data.get('sources', [])
        for i, source in enumerate(sources, start=1):
            doc.add_paragraph(f"{i}. {source}", style='List Number')
        
        # Сохранение
        doc.save(output_path)
        logger.info(f"Синопсис сохранен: {output_path}")
        
        return output_path
    
    @staticmethod
    def generate_json(data: dict, output_path: str = "synopsis.json") -> str:
        """
        Генерация JSON синопсиса
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"JSON синопсис сохранен: {output_path}")
        return output_path
    
    @staticmethod
    def generate_markdown(data: dict, output_path: str = "synopsis.md") -> str:
        """
        Генерация Markdown синопсиса
        """
        md_content = f"""# СИНОПСИС ПРОТОКОЛА

## Исследование биоэквивалентности

**Препарат:** {data.get('inn', 'N/A')}  
**Дата:** {datetime.now().strftime('%d.%m.%Y')}

---

## 1. ОБОСНОВАНИЕ ДИЗАЙНА ИССЛЕДОВАНИЯ

{data.get('design_rationale', 'N/A')}

## 2. ЦЕЛИ И ЗАДАЧИ

Цель: Оценка биоэквивалентности исследуемого препарата и референтного препарата.

## 3. ДИЗАЙН ИССЛЕДОВАНИЯ

- **Тип дизайна:** {data.get('study_design', {}).get('design', 'N/A')}
- **Количество периодов:** {data.get('study_design', {}).get('periods', 'N/A')}
- **Washout период:** {data.get('study_design', {}).get('washout_duration', 'N/A')}

## 4. РАЗМЕР ВЫБОРКИ

- **CVintra:** {data.get('sample_size', {}).get('cvintra', 'N/A')}%
- **Базовый размер:** {data.get('sample_size', {}).get('base_sample_size', 'N/A')}
- **Итоговый размер:** {data.get('sample_size', {}).get('final_sample_size', 'N/A')}

### Расчет:
"""
        
        for step in data.get('sample_size', {}).get('steps', []):
            md_content += f"\n{step}"
        
        md_content += """

## 5. ФАРМАКОКИНЕТИЧЕСКИЕ ПАРАМЕТРЫ

| Параметр | Описание | Единица |
|----------|----------|---------|
| Cmax | Максимальная концентрация | нг/мл |
| AUC | Площадь под кривой | нг·ч/мл |
| Tmax | Время достижения Cmax | ч |
| T½ | Период полувыведения | ч |
| CVintra | Внутрисубъектная вариабельность | % |

## 6. СТАТИСТИЧЕСКАЯ МЕТОДОЛОГИЯ

90% доверительный интервал для отношения геометрических средних Cmax и AUC должен находиться в пределах 80.00% - 125.00%.

## 7. БИБЛИОГРАФИЯ

"""
        
        for i, source in enumerate(data.get('sources', []), start=1):
            md_content += f"\n{i}. {source}"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        logger.info(f"Markdown синопсис сохранен: {output_path}")
        return output_path